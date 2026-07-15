from pathlib import Path

from docx import Document

from bidpilot_data.chunking.chunk import chunk_documents, estimate_tokens, secondary_token_split
from bidpilot_data.cleaning.clean import clean_text
from bidpilot_data.parsers.parse import parse_docx, parse_plain_text
from bidpilot_data.schemas import DocumentRecord, DocumentType, ParseStatus
from bidpilot_data.utils import write_jsonl


def test_clean_preserves_money_and_codes():
    text = "预算金额：3,500,000元\n项目编号：DEMO-2026-001\n保证金 5%\n\n\n重复\n重复"
    out = clean_text(text)
    assert "3,500,000" in out
    assert "DEMO-2026-001" in out
    assert "5%" in out


def test_parse_docx_and_plaintext(tmp_path: Path):
    docx_path = tmp_path / "a.docx"
    d = Document()
    d.add_heading("第一章 资格要求", level=1)
    d.add_paragraph("投标人须提供有效营业执照。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "评分项"
    table.cell(0, 1).text = "分值"
    table.cell(1, 0).text = "业绩"
    table.cell(1, 1).text = "15"
    d.save(docx_path)
    pages, method, status = parse_docx(docx_path, "doc-1")
    assert status == ParseStatus.success
    assert "营业执照" in pages[0].text
    assert pages[0].tables

    txt = tmp_path / "b.txt"
    txt.write_text("第1条 投标无效条款。\n投标人不得虚假响应，否则投标无效。", encoding="utf-8")
    pages2, _, status2 = parse_plain_text(txt, "doc-2")
    assert status2 == ParseStatus.success
    assert pages2[0].page_number == 1


def test_chunk_stability(tmp_datasets: Path, tmp_repo: Path):
    doc_id = "11111111-1111-1111-1111-111111111111"
    project_id = "22222222-2222-2222-2222-222222222222"
    write_jsonl(
        tmp_datasets / "manifests" / "documents.jsonl",
        [
            DocumentRecord(
                document_id=doc_id,
                project_id=project_id,
                original_filename="t.txt",
                sha256="a" * 64,
                file_size=10,
                storage_path="raw/documents/t.txt",
                parse_status=ParseStatus.success,
                document_type=DocumentType.tender,
                source_url="file:///tmp/t.txt",
            )
        ],
    )
    text = "\n".join(
        [
            "第一章 资格要求",
            "1.1 投标人须提供有效营业执照。",
            "1.2 投标人应当具备相关资质，否则投标无效。",
            "第二章 技术要求",
            "2.1 系统应支持不少于200路视频接入。",
        ]
    )
    (tmp_datasets / "interim" / "cleaned" / f"{doc_id}.jsonl").write_text(
        '{"document_id":"%s","page_number":1,"text":%s,"tables":[],"headings":[],"ocr_used":false}\n'
        % (doc_id, __import__("json").dumps(text)),
        encoding="utf-8",
    )
    s1 = chunk_documents(resume=False)
    s2 = chunk_documents(resume=False)
    assert s1["chunks"] == s2["chunks"]
    assert s1["chunks"] >= 1
    parts = secondary_token_split("句号。" * 500, max_tokens=50, overlap=5)
    assert len(parts) > 1
    assert estimate_tokens("汉字abc") >= 3
