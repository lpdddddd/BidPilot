from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup

from bidpilot_data.logging import get_logger, log_stats
from bidpilot_data.schemas import DocumentRecord, ParseStatus, ParsedPage
from bidpilot_data.settings import get_settings
from bidpilot_data.utils import CheckpointStore, ensure_dir, read_jsonl, write_json, write_jsonl

log = get_logger(__name__)

HEADING_RE = re.compile(r"^(第[一二三四五六七八九十百千0-9]+[章节条款项]|[0-9]+(\.[0-9]+){0,3}|（[一二三四五六七八九十]+）)")


def _is_probably_scanned(pages_text: list[str]) -> bool:
    if not pages_text:
        return True
    chars = sum(len(t.strip()) for t in pages_text)
    avg = chars / max(len(pages_text), 1)
    return avg < 40


def parse_pdf(path: Path, document_id: str) -> tuple[list[ParsedPage], str, ParseStatus]:
    try:
        import fitz  # pymupdf
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pymupdf is required for PDF parsing") from exc

    doc = fitz.open(path)
    pages: list[ParsedPage] = []
    texts: list[str] = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        texts.append(text)
        headings = [ln.strip() for ln in text.splitlines() if HEADING_RE.match(ln.strip())][:20]
        pages.append(
            ParsedPage(
                document_id=document_id,
                page_number=i + 1,
                text=text,
                tables=[],
                headings=headings,
                bbox_metadata=None,
                ocr_used=False,
            )
        )
    doc.close()

    if _is_probably_scanned(texts):
        # Optional OCR
        try:
            import paddleocr  # noqa: F401

            ocr_available = True
        except Exception:  # noqa: BLE001
            ocr_available = False
        if not ocr_available:
            return pages, "pymupdf_text", ParseStatus.ocr_required
        # OCR path intentionally minimal: mark partial if installed but not fully wired in CI.
        return pages, "pymupdf_text+ocr_pending", ParseStatus.partial
    status = ParseStatus.success if any(p.text.strip() for p in pages) else ParseStatus.failed
    return pages, "pymupdf_text", status


def parse_docx(path: Path, document_id: str) -> tuple[list[ParsedPage], str, ParseStatus]:
    from docx import Document

    document = Document(str(path))
    paragraphs: list[str] = []
    headings: list[str] = []
    tables: list[list[list[str]]] = []
    for p in document.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        paragraphs.append(text)
        if HEADING_RE.match(text) or (p.style and "Heading" in str(p.style.name)):
            headings.append(text)
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
            # Preserve table numeric / qualification content in linear text
            paragraphs.append("\n".join(["\t".join(r) for r in rows]))

    page = ParsedPage(
        document_id=document_id,
        page_number=1,
        text="\n".join(paragraphs),
        tables=tables,
        headings=headings,
        ocr_used=False,
    )
    status = ParseStatus.success if page.text.strip() else ParseStatus.failed
    return [page], "python-docx", status


def parse_html(path: Path, document_id: str) -> tuple[list[ParsedPage], str, ParseStatus]:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("\n")
    headings = [h.get_text(strip=True) for h in soup.find_all(re.compile("^h[1-6]$"))]
    page = ParsedPage(
        document_id=document_id,
        page_number=1,
        text=text,
        headings=headings,
        ocr_used=False,
    )
    return [page], "beautifulsoup", ParseStatus.success if text.strip() else ParseStatus.failed


def parse_plain_text(path: Path, document_id: str) -> tuple[list[ParsedPage], str, ParseStatus]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    # Split into pseudo-pages by form feed or every ~3000 chars keeping section boundaries.
    chunks = [c for c in text.split("\f") if c.strip()] or [text]
    pages = [
        ParsedPage(
            document_id=document_id,
            page_number=i + 1,
            text=chunk,
            headings=[ln.strip() for ln in chunk.splitlines() if HEADING_RE.match(ln.strip())][:20],
            ocr_used=False,
        )
        for i, chunk in enumerate(chunks)
    ]
    return pages, "plaintext", ParseStatus.success


def parse_one(path: Path, document_id: str) -> tuple[list[ParsedPage], str, ParseStatus]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path, document_id)
    if suffix in {".docx"}:
        return parse_docx(path, document_id)
    if suffix in {".html", ".htm"}:
        return parse_html(path, document_id)
    if suffix in {".txt", ".md", ".json"}:
        return parse_plain_text(path, document_id)
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return (
            [
                ParsedPage(
                    document_id=document_id,
                    page_number=1,
                    text="",
                    ocr_used=False,
                )
            ],
            "image",
            ParseStatus.ocr_required,
        )
    return (
        [ParsedPage(document_id=document_id, page_number=1, text="", ocr_used=False)],
        "unsupported",
        ParseStatus.failed,
    )


def parse_documents(*, resume: bool = True, dry_run: bool = False) -> dict[str, Any]:
    settings = get_settings()
    docs_path = settings.datasets_root / "manifests" / "documents.jsonl"
    docs = [DocumentRecord.model_validate(d) for d in read_jsonl(docs_path)]
    out_dir = ensure_dir(settings.datasets_root / "interim" / "parsed")
    ckpt = CheckpointStore(settings.datasets_root / "reports" / "checkpoints" / "parse.json")
    stats = {"seen": 0, "success": 0, "partial": 0, "ocr_required": 0, "failed": 0, "skipped": 0}

    updated: list[dict[str, Any]] = []
    for doc in docs:
        stats["seen"] += 1
        if resume and ckpt.done(doc.document_id):
            stats["skipped"] += 1
            updated.append(doc.model_dump(mode="json"))
            continue
        path = settings.datasets_root / doc.storage_path
        if dry_run:
            updated.append(doc.model_dump(mode="json"))
            continue
        if not path.exists():
            doc.parse_status = ParseStatus.failed
            stats["failed"] += 1
            ckpt.failed(doc.document_id, "missing file")
            updated.append(doc.model_dump(mode="json"))
            continue
        try:
            pages, method, status = parse_one(path, doc.document_id)
            doc.parse_method = method
            doc.parse_status = status
            doc.page_count = len(pages)
            write_jsonl(out_dir / f"{doc.document_id}.jsonl", pages)
            write_json(
                out_dir / f"{doc.document_id}.meta.json",
                {"document_id": doc.document_id, "parse_method": method, "parse_status": status, "page_count": len(pages)},
            )
            stats[status.value if status.value in stats else "failed"] = stats.get(status.value, 0) + (
                1 if status.value in stats else 0
            )
            if status == ParseStatus.success:
                stats["success"] += 1
            elif status == ParseStatus.partial:
                stats["partial"] += 1
            elif status == ParseStatus.ocr_required:
                stats["ocr_required"] += 1
            else:
                stats["failed"] += 1
            ckpt.mark_done(doc.document_id, {"status": status.value})
        except Exception as exc:  # noqa: BLE001
            doc.parse_status = ParseStatus.failed
            stats["failed"] += 1
            ckpt.failed(doc.document_id, str(exc))
            log.warning("parse failed document_id=%s err=%s", doc.document_id, exc)
        updated.append(doc.model_dump(mode="json"))

    write_jsonl(docs_path, updated)
    log_stats(log, "parse", stats)
    return stats
