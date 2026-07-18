"""Text extraction for uploaded tender documents.

Supported: PDF (pypdf), DOCX (python-docx), TXT, HTML/HTM (BeautifulSoup),
XLSX (openpyxl). Anything else, plus scanned PDFs without an extractable text
layer, is reported honestly as failed / ocr_required. No fake results.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from app.models.enums import ParseStatus

PARSER_NAME = "bidpilot-basic-parser"
PARSER_VERSION = "1.1.0"

# A PDF whose pages yield fewer characters than this on average is treated as
# a scanned document that needs OCR.
_MIN_CHARS_PER_PAGE = 10


@dataclass
class PageSpan:
    """Character range of one PDF page inside the extracted text."""

    page: int
    char_start: int
    char_end: int


@dataclass
class ParseResult:
    status: ParseStatus
    text: str = ""
    page_count: int | None = None
    error: str | None = None
    # Only populated for PDF: real page-to-character mapping. None means the
    # format has no reliable page notion (DOCX/TXT/HTML/XLSX).
    page_spans: list[PageSpan] | None = None

    @property
    def extracted_characters(self) -> int:
        return len(self.text)


def parse_document(content: bytes, extension: str) -> ParseResult:
    ext = extension.lower().lstrip(".")
    try:
        if ext == "pdf":
            return _parse_pdf(content)
        if ext == "docx":
            return _parse_docx(content)
        if ext == "txt":
            return _parse_txt(content)
        if ext in {"html", "htm"}:
            return _parse_html(content)
        if ext == "xlsx":
            return _parse_xlsx(content)
        return ParseResult(
            status=ParseStatus.failed,
            error=f"不支持解析的文件类型: .{ext}（DOC/WPS 等格式请先转换为 PDF 或 DOCX）",
        )
    except Exception as exc:  # noqa: BLE001 - parse failures must not crash the task
        return ParseResult(status=ParseStatus.failed, error=f"解析异常: {exc}")


def _parse_pdf(content: bytes) -> ParseResult:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            return ParseResult(status=ParseStatus.failed, error="PDF 已加密，无法提取文本")

    page_count = len(reader.pages)
    pages_text: list[str] = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - keep going; a single bad page is not fatal
            pages_text.append("")

    # Build the extracted text and the page-to-character mapping together so
    # every span is exact with respect to the final artifact.
    segments: list[str] = []
    spans: list[PageSpan] = []
    cursor = 0
    for page_number, raw in enumerate(pages_text, start=1):
        part = raw.strip()
        if not part:
            continue
        if segments:
            cursor += 1  # "\n" separator between page segments
        start = cursor
        segments.append(part)
        cursor += len(part)
        spans.append(PageSpan(page=page_number, char_start=start, char_end=cursor))
    text = "\n".join(segments)

    if page_count > 0 and len(text) < _MIN_CHARS_PER_PAGE * page_count:
        return ParseResult(
            status=ParseStatus.ocr_required,
            text=text,
            page_count=page_count,
            error="PDF 无可提取文本层，疑似扫描件，需要 OCR",
        )
    return ParseResult(
        status=ParseStatus.success,
        text=text,
        page_count=page_count,
        page_spans=spans,
    )


def _parse_docx(content: bytes) -> ParseResult:
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(content))
    parts: list[str] = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        return ParseResult(status=ParseStatus.failed, error="DOCX 中未提取到任何文本")
    return ParseResult(status=ParseStatus.success, text=text)


_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _looks_like_text(text: str) -> bool:
    """Reject decodes that are structurally not natural text (binary sniff)."""
    if not text.strip():
        return False
    control_count = len(_CONTROL_CHARS_RE.findall(text))
    return control_count / len(text) < 0.02


def _parse_txt(content: bytes) -> ParseResult:
    # Binary files renamed to .txt must fail honestly, not "parse" via a
    # lenient codec. NUL bytes are a hard binary signal.
    if b"\x00" in content:
        return ParseResult(status=ParseStatus.failed, error="文件包含二进制内容，不是文本文件")

    candidates = ["utf-8", "gb18030"]
    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        candidates = ["utf-16", *candidates]

    for encoding in candidates:
        try:
            text = content.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
        if not text.strip():
            return ParseResult(status=ParseStatus.failed, error="文本文件内容为空")
        if _looks_like_text(text):
            return ParseResult(status=ParseStatus.success, text=text)
    return ParseResult(
        status=ParseStatus.failed,
        error="无法识别文本编码或内容不是可读文本",
    )


def _parse_html(content: bytes) -> ParseResult:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "noscript", "template", "head"]):
        tag.decompose()
    text = "\n".join(
        line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip()
    )
    if not text:
        return ParseResult(status=ParseStatus.failed, error="HTML 中未提取到可读正文")
    return ParseResult(status=ParseStatus.success, text=text)


def _parse_xlsx(content: bytes) -> ParseResult:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            parts.append(f"[工作表] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value).strip() for value in row if value is not None]
                if cells:
                    parts.append("\t".join(cells))
    finally:
        workbook.close()
    text = "\n".join(parts)
    has_cell_content = any(not part.startswith("[工作表] ") for part in parts)
    if not has_cell_content:
        return ParseResult(status=ParseStatus.failed, error="XLSX 中未提取到任何单元格文本")
    return ParseResult(status=ParseStatus.success, text=text)
