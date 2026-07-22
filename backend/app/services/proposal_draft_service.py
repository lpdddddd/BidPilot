"""Auditable proposal drafting workspace service."""

from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from datetime import UTC, datetime
from typing import Any
from uuid import UUID

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from app.models.document import DocumentChunk
from app.models.enums import (
    ActorAuthn,
    EvidenceMatchStatus,
    ExtractionRunStatus,
    MatchReviewStatus,
    ProposalDraftReviewAction,
    ProposalDraftSourceRole,
    ProposalDraftStatus,
    ProposalDraftVersionKind,
)
from app.models.match_run import RequirementEvidenceMatch, RequirementEvidenceMatchLink
from app.models.project import BidProject
from app.models.proposal_draft import (
    ProposalDraft,
    ProposalDraftGenerationRun,
    ProposalDraftReview,
    ProposalDraftSource,
    ProposalDraftVersion,
)
from app.models.requirement import EvidenceLink, Requirement
from app.schemas.proposal_draft import (
    DISCLAIMER,
    EligibilityRequirementItem,
    ProposalDraftCreateRequest,
    ProposalDraftDetail,
    ProposalDraftEligibilityResponse,
    ProposalDraftListResponse,
    ProposalDraftManualRevisionRequest,
    ProposalDraftReopenRequest,
    ProposalDraftReviewRead,
    ProposalDraftReviewRequest,
    ProposalDraftRunResponse,
    ProposalDraftSourceRead,
    ProposalDraftSummary,
    ProposalDraftVersionDetail,
    ProposalDraftVersionListResponse,
    ProposalDraftVersionSummary,
)
from app.services.llm_client import LlmClient, LlmError, get_llm_client
from app.services.proposal_draft_validate import (
    CitationMeta,
    DraftValidationError,
    WhitelistContext,
    content_has_unevidenced_manual,
    normalize_text,
    render_markdown,
    validate_structured_content,
)

logger = logging.getLogger("bidpilot.proposal_draft")

POSITIVE_STATUSES = frozenset(
    {EvidenceMatchStatus.supported, EvidenceMatchStatus.partially_supported}
)

_SYSTEM_PROMPT = """你是投标响应准备助手。只能基于服务端提供的白名单事实输出结构化 JSON。
禁止输出投标结论、报价、工期承诺、法律承诺、盖章承诺或自由散文。
禁止编造 citation_ids / source_quote_ids / 页码 / section / clause。
只返回 JSON 对象，字段：title, sections, compliance_matrix, warnings。
block_kind 必须与给定 match_status 对应：
supported→supported_response；partially_supported→partial_response；
insufficient_evidence→material_gap；conflicting_evidence→risk_item；
not_applicable→scope_item。
事实性 block 必须引用白名单 citation_ids 与 source_quote_ids。
"""


def _payload_hash(payload: dict[str, Any]) -> str:
    raw = json.dumps(payload, sort_keys=True, ensure_ascii=False, default=str)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _parse_llm_json(raw: str) -> dict[str, Any]:
    text = (raw or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise DraftValidationError(f"invalid LLM JSON: {exc}") from exc
    if not isinstance(data, dict):
        raise DraftValidationError("LLM JSON must be an object")
    return data


def _location_from_chunk(
    *,
    document_file_name: str | None,
    document_type: str | None,
    chunk: DocumentChunk | None,
    document_id: UUID | None,
    chunk_id: UUID | None,
    project_id: UUID,
) -> dict[str, Any]:
    loc: dict[str, Any] = {
        "document_id": str(document_id) if document_id else None,
        "chunk_id": str(chunk_id) if chunk_id else None,
        "document_file_name": document_file_name,
        "document_type": document_type,
    }
    if chunk is not None:
        loc.update(
            {
                "chunk_index": chunk.chunk_index,
                "section": chunk.section,
                "clause_id": chunk.clause_id,
                "page_start": chunk.page_start,
                "page_end": chunk.page_end,
            }
        )
    if document_id and chunk_id:
        loc["document_center_path"] = (
            f"/projects/{project_id}?tab=documents&documentId={document_id}"
            f"&chunkId={chunk_id}"
        )
    return loc


class ProposalDraftService:
    def __init__(self, db: Session, llm: LlmClient | None = None) -> None:
        self.db = db
        self.llm = llm if llm is not None else get_llm_client()

    # ------------------------------------------------------------------ project

    def _require_project(self, project_id: UUID) -> BidProject:
        project = self.db.get(BidProject, project_id)
        if project is None:
            raise HTTPException(status_code=404, detail="项目不存在")
        return project

    def _lock_draft(self, project_id: UUID, draft_id: UUID) -> ProposalDraft:
        draft = self.db.execute(
            select(ProposalDraft)
            .where(
                ProposalDraft.id == draft_id,
                ProposalDraft.project_id == project_id,
            )
            .with_for_update()
        ).scalar_one_or_none()
        if draft is None:
            raise HTTPException(status_code=404, detail="草稿不存在")
        return draft

    # -------------------------------------------------------------- eligibility

    def eligibility(
        self, project_id: UUID, requirement_ids: list[UUID] | None = None
    ) -> ProposalDraftEligibilityResponse:
        self._require_project(project_id)
        reqs = list(
            self.db.scalars(
                select(Requirement).where(Requirement.project_id == project_id)
            )
        )
        if requirement_ids:
            wanted = set(requirement_ids)
            reqs = [r for r in reqs if r.id in wanted]

        matches = list(
            self.db.scalars(
                select(RequirementEvidenceMatch)
                .where(
                    RequirementEvidenceMatch.project_id == project_id,
                    RequirementEvidenceMatch.lifecycle_status == "active",
                )
                .options(selectinload(RequirementEvidenceMatch.requirement))
            )
        )
        by_req = {m.requirement_id: m for m in matches}

        eligible: list[EligibilityRequirementItem] = []
        excluded: list[EligibilityRequirementItem] = []
        material_gaps: list[EligibilityRequirementItem] = []
        risks: list[EligibilityRequirementItem] = []
        scope_items: list[EligibilityRequirementItem] = []

        for req in reqs:
            match = by_req.get(req.id)
            item = self._classify_requirement(req, match)
            bucket = {
                "positive": eligible,
                "material_gap": material_gaps,
                "risk": risks,
                "scope": scope_items,
                "excluded": excluded,
                "no_match": excluded,
            }[item.eligibility]
            bucket.append(item)

        return ProposalDraftEligibilityResponse(
            project_id=project_id,
            eligible=eligible,
            excluded=excluded,
            material_gaps=material_gaps,
            risks=risks,
            scope_items=scope_items,
        )

    def _classify_requirement(
        self, req: Requirement, match: RequirementEvidenceMatch | None
    ) -> EligibilityRequirementItem:
        if match is None:
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                eligibility="no_match",
                reason="无 active Match",
                draft_handling="excluded_from_positive_body",
            )
        if match.review_status != MatchReviewStatus.confirmed:
            reason_map = {
                MatchReviewStatus.pending: "pending_review",
                MatchReviewStatus.rejected: "rejected_match",
                MatchReviewStatus.needs_more_material: "needs_more_material",
            }
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                match_id=match.id,
                match_status=match.status,
                review_status=match.review_status,
                eligibility="excluded",
                reason=reason_map.get(match.review_status, match.review_status.value),
                draft_handling="listed_as_pending_human_action_only",
            )
        if match.status in POSITIVE_STATUSES:
            handling = (
                "supported_response"
                if match.status == EvidenceMatchStatus.supported
                else "partial_response_with_gap"
            )
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                match_id=match.id,
                match_status=match.status,
                review_status=match.review_status,
                eligibility="positive",
                reason="confirmed_active_positive",
                draft_handling=handling,
            )
        if match.status == EvidenceMatchStatus.insufficient_evidence:
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                match_id=match.id,
                match_status=match.status,
                review_status=match.review_status,
                eligibility="material_gap",
                reason="confirmed_insufficient_evidence",
                draft_handling="material_gap_list_only",
            )
        if match.status == EvidenceMatchStatus.conflicting_evidence:
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                match_id=match.id,
                match_status=match.status,
                review_status=match.review_status,
                eligibility="risk",
                reason="confirmed_conflicting_evidence",
                draft_handling="risk_list_with_dual_evidence",
            )
        if match.status == EvidenceMatchStatus.not_applicable:
            return EligibilityRequirementItem(
                requirement_id=req.id,
                title=req.title,
                category=req.category,
                match_id=match.id,
                match_status=match.status,
                review_status=match.review_status,
                eligibility="scope",
                reason="confirmed_not_applicable",
                draft_handling="scope_list_with_dual_evidence",
            )
        return EligibilityRequirementItem(
            requirement_id=req.id,
            title=req.title,
            category=req.category,
            match_id=match.id,
            match_status=match.status,
            review_status=match.review_status,
            eligibility="excluded",
            reason=f"unsupported_status:{match.status}",
            draft_handling="excluded_from_positive_body",
        )

    # ------------------------------------------------------------------- CRUD

    def list_drafts(self, project_id: UUID) -> ProposalDraftListResponse:
        self._require_project(project_id)
        rows = list(
            self.db.scalars(
                select(ProposalDraft)
                .where(ProposalDraft.project_id == project_id)
                .order_by(ProposalDraft.updated_at.desc())
            )
        )
        items = [self._to_summary(d) for d in rows]
        return ProposalDraftListResponse(items=items, total=len(items))

    def get_draft(self, project_id: UUID, draft_id: UUID) -> ProposalDraftDetail:
        self._require_project(project_id)
        draft = (
            self.db.execute(
                select(ProposalDraft)
                .where(
                    ProposalDraft.id == draft_id,
                    ProposalDraft.project_id == project_id,
                )
                .options(
                    selectinload(ProposalDraft.reviews),
                    selectinload(ProposalDraft.generation_runs),
                )
            )
            .scalar_one_or_none()
        )
        if draft is None:
            raise HTTPException(status_code=404, detail="草稿不存在")
        summary = self._to_summary(draft)
        current = None
        if draft.current_version_id:
            current = self._version_detail(project_id, draft.current_version_id)
        reviews = [
            ProposalDraftReviewRead.model_validate(r) for r in (draft.reviews or [])[:20]
        ]
        latest_run = None
        runs = sorted(
            draft.generation_runs or [],
            key=lambda r: r.created_at,
            reverse=True,
        )
        if runs:
            latest_run = ProposalDraftRunResponse.model_validate(runs[0])
        return ProposalDraftDetail(
            **summary.model_dump(),
            current_version=current,
            recent_reviews=reviews,
            latest_run=latest_run,
        )

    def list_versions(
        self, project_id: UUID, draft_id: UUID
    ) -> ProposalDraftVersionListResponse:
        self._require_project(project_id)
        draft = self.db.get(ProposalDraft, draft_id)
        if draft is None or draft.project_id != project_id:
            raise HTTPException(status_code=404, detail="草稿不存在")
        rows = list(
            self.db.scalars(
                select(ProposalDraftVersion)
                .where(
                    ProposalDraftVersion.draft_id == draft_id,
                    ProposalDraftVersion.project_id == project_id,
                )
                .order_by(ProposalDraftVersion.version_number.desc())
            )
        )
        items = [self._version_summary(v) for v in rows]
        return ProposalDraftVersionListResponse(items=items, total=len(items))

    def get_version(
        self, project_id: UUID, draft_id: UUID, version_id: UUID
    ) -> ProposalDraftVersionDetail:
        self._require_project(project_id)
        version = self.db.execute(
            select(ProposalDraftVersion)
            .where(
                ProposalDraftVersion.id == version_id,
                ProposalDraftVersion.draft_id == draft_id,
                ProposalDraftVersion.project_id == project_id,
            )
            .options(selectinload(ProposalDraftVersion.sources))
        ).scalar_one_or_none()
        if version is None:
            raise HTTPException(status_code=404, detail="版本不存在")
        return self._version_detail_from_model(version)

    # -------------------------------------------------------------- generation

    def start_generation(
        self,
        project_id: UUID,
        payload: ProposalDraftCreateRequest,
        *,
        idempotency_key: str | None = None,
    ) -> ProposalDraftRunResponse:
        self._require_project(project_id)
        req_ids = list(dict.fromkeys(payload.requirement_ids))
        if not req_ids:
            raise HTTPException(status_code=422, detail="requirement_ids 不能为空")

        if idempotency_key:
            existing = self.db.execute(
                select(ProposalDraftGenerationRun).where(
                    ProposalDraftGenerationRun.project_id == project_id,
                    ProposalDraftGenerationRun.idempotency_key == idempotency_key,
                )
            ).scalar_one_or_none()
            if existing is not None:
                expected = _payload_hash(
                    {
                        "title": payload.title,
                        "requirement_ids": [str(x) for x in req_ids],
                        "mode": payload.mode.value,
                    }
                )
                stored = (existing.config_json or {}).get("payload_hash")
                if stored and stored != expected:
                    raise HTTPException(
                        status_code=409,
                        detail="Idempotency-Key 已用于不同请求载荷",
                    )
                return ProposalDraftRunResponse.model_validate(existing)

        eligibility = self.eligibility(project_id, req_ids)
        eligible_count = len(eligibility.eligible) + len(eligibility.material_gaps) + len(
            eligibility.risks
        ) + len(eligibility.scope_items)
        excluded_count = len(eligibility.excluded)
        excluded_reasons = [
            f"{x.requirement_id}:{x.reason}" for x in eligibility.excluded[:20]
        ]

        # At least one confirmed active match of any kind, or fail early
        confirmed_any = eligible_count > 0
        if not confirmed_any:
            raise HTTPException(
                status_code=422,
                detail="所选 Requirement 无已确认 active Match，无法生成草稿",
            )

        run = ProposalDraftGenerationRun(
            project_id=project_id,
            status=ExtractionRunStatus.queued,
            mode=payload.mode,
            title=payload.title,
            requested_requirement_ids=[str(x) for x in req_ids],
            eligible_requirement_count=eligible_count,
            excluded_requirement_count=excluded_count,
            excluded_reason_summary="; ".join(excluded_reasons) or None,
            created_by=payload.created_by,
            idempotency_key=idempotency_key,
            config_json={
                "payload_hash": _payload_hash(
                    {
                        "title": payload.title,
                        "requirement_ids": [str(x) for x in req_ids],
                        "mode": payload.mode.value,
                    }
                ),
                "cancel_requested": False,
                "model": getattr(self.llm, "model", None),
            },
        )
        self.db.add(run)
        self.db.commit()
        self.db.refresh(run)
        return ProposalDraftRunResponse.model_validate(run)

    def get_run(self, project_id: UUID, run_id: UUID) -> ProposalDraftRunResponse:
        self._require_project(project_id)
        run = self.db.get(ProposalDraftGenerationRun, run_id)
        if run is None or run.project_id != project_id:
            raise HTTPException(status_code=404, detail="生成任务不存在")
        return ProposalDraftRunResponse.model_validate(run)

    def cancel_run(self, project_id: UUID, run_id: UUID) -> ProposalDraftRunResponse:
        self._require_project(project_id)
        run = self.db.execute(
            select(ProposalDraftGenerationRun)
            .where(
                ProposalDraftGenerationRun.id == run_id,
                ProposalDraftGenerationRun.project_id == project_id,
            )
            .with_for_update()
        ).scalar_one_or_none()
        if run is None:
            raise HTTPException(status_code=404, detail="生成任务不存在")
        if run.status in (
            ExtractionRunStatus.succeeded,
            ExtractionRunStatus.failed,
            ExtractionRunStatus.cancelled,
        ):
            raise HTTPException(status_code=409, detail="任务已结束，无法取消")
        now = datetime.now(UTC)
        cfg = dict(run.config_json or {})
        cfg["cancel_requested"] = True
        run.config_json = cfg
        run.cancel_requested_at = now
        run.status = ExtractionRunStatus.cancelled
        run.finished_at = now
        run.error_summary = "任务已取消，未写入草稿版本"
        self.db.commit()
        self.db.refresh(run)
        return ProposalDraftRunResponse.model_validate(run)

    def _is_cancel_requested(self, run_id: UUID) -> bool:
        run = self.db.get(ProposalDraftGenerationRun, run_id)
        if run is None:
            return True
        self.db.refresh(run)
        if run.status == ExtractionRunStatus.cancelled:
            return True
        return bool((run.config_json or {}).get("cancel_requested"))

    def execute_run(self, run_id: UUID) -> None:
        run = self.db.get(ProposalDraftGenerationRun, run_id)
        if run is None:
            logger.warning("Draft run %s missing", run_id)
            return
        if run.status in (
            ExtractionRunStatus.succeeded,
            ExtractionRunStatus.failed,
            ExtractionRunStatus.cancelled,
        ):
            return
        if self._is_cancel_requested(run_id):
            self._mark_cancelled(run)
            return

        run.status = ExtractionRunStatus.running
        run.started_at = datetime.now(UTC)
        run.error_summary = None
        self.db.commit()

        try:
            if self._is_cancel_requested(run_id):
                self._mark_cancelled(run)
                return

            req_ids = [UUID(str(x)) for x in (run.requested_requirement_ids or [])]
            whitelist, source_rows = self._build_whitelist(run.project_id, req_ids)
            if self._is_cancel_requested(run_id):
                self._mark_cancelled(run)
                return

            llm_payload = self._call_llm(run, whitelist)
            if self._is_cancel_requested(run_id):
                self._mark_cancelled(run)
                return

            try:
                content = validate_structured_content(llm_payload, whitelist)
            except DraftValidationError as exc:
                self._mark_failed(run, f"结构化校验失败: {exc.message}")
                return

            markdown = render_markdown(content)
            snapshot_hash = hashlib.sha256(
                json.dumps(
                    {
                        "sources": [
                            {
                                "requirement_id": str(s["requirement_id"]),
                                "match_id": str(s["match_id"]) if s.get("match_id") else None,
                                "evidence_link_id": str(s["evidence_link_id"])
                                if s.get("evidence_link_id")
                                else None,
                                "source_role": s["source_role"],
                                "source_quote": s.get("source_quote"),
                            }
                            for s in source_rows
                        ]
                    },
                    sort_keys=True,
                    ensure_ascii=False,
                ).encode("utf-8")
            ).hexdigest()

            # Stage in memory then atomic persist with FOR UPDATE on run
            self._persist_success(
                run,
                content=content,
                markdown=markdown,
                source_rows=source_rows,
                snapshot_hash=snapshot_hash,
            )
        except LlmError as exc:
            self._mark_failed(run, f"LLM 失败: {exc.message}")
        except DraftValidationError as exc:
            self._mark_failed(run, f"校验失败: {exc.message}")
        except Exception as exc:  # noqa: BLE001
            logger.exception("Draft run %s failed", run_id)
            self.db.rollback()
            run = self.db.get(ProposalDraftGenerationRun, run_id)
            if run is not None:
                self._mark_failed(run, f"内部错误: {exc}")

    def _mark_cancelled(self, run: ProposalDraftGenerationRun) -> None:
        run.status = ExtractionRunStatus.cancelled
        run.finished_at = datetime.now(UTC)
        run.cancel_requested_at = run.cancel_requested_at or datetime.now(UTC)
        run.error_summary = "任务已取消，未写入草稿版本"
        cfg = dict(run.config_json or {})
        cfg["cancel_requested"] = True
        run.config_json = cfg
        self.db.commit()

    def _mark_failed(self, run: ProposalDraftGenerationRun, reason: str) -> None:
        run.status = ExtractionRunStatus.failed
        run.finished_at = datetime.now(UTC)
        run.error_summary = reason[:2000]
        self.db.commit()

    def _persist_success(
        self,
        run: ProposalDraftGenerationRun,
        *,
        content: dict[str, Any],
        markdown: str,
        source_rows: list[dict[str, Any]],
        snapshot_hash: str,
    ) -> None:
        locked = self.db.execute(
            select(ProposalDraftGenerationRun)
            .where(ProposalDraftGenerationRun.id == run.id)
            .with_for_update()
        ).scalar_one()
        if locked.status == ExtractionRunStatus.cancelled or bool(
            (locked.config_json or {}).get("cancel_requested")
        ):
            locked.status = ExtractionRunStatus.cancelled
            locked.finished_at = datetime.now(UTC)
            locked.error_summary = "任务已取消，未写入草稿版本"
            self.db.commit()
            return

        draft = ProposalDraft(
            project_id=locked.project_id,
            title=locked.title,
            status=ProposalDraftStatus.draft_pending_review,
            created_by=locked.created_by,
            review_lock_version=0,
            metadata_json={
                "mode": locked.mode.value,
                "disclaimer": DISCLAIMER,
            },
        )
        self.db.add(draft)
        self.db.flush()

        version = ProposalDraftVersion(
            project_id=locked.project_id,
            draft_id=draft.id,
            parent_version_id=None,
            version_number=1,
            version_kind=ProposalDraftVersionKind.generated,
            generation_run_id=locked.id,
            source_snapshot_hash=snapshot_hash,
            content_json=content,
            content_markdown=markdown,
            created_by=locked.created_by,
            is_current=True,
        )
        self.db.add(version)
        self.db.flush()

        for row in source_rows:
            self.db.add(
                ProposalDraftSource(
                    project_id=locked.project_id,
                    draft_version_id=version.id,
                    requirement_id=row.get("requirement_id"),
                    match_id=row.get("match_id"),
                    evidence_link_id=row.get("evidence_link_id"),
                    source_role=ProposalDraftSourceRole(row["source_role"]),
                    source_quote=row.get("source_quote"),
                    location_json=row.get("location_json"),
                )
            )

        draft.current_version_id = version.id
        locked.draft_id = draft.id
        locked.draft_version_id = version.id
        locked.status = ExtractionRunStatus.succeeded
        locked.finished_at = datetime.now(UTC)
        locked.error_summary = None
        self.db.commit()

    # ----------------------------------------------------------- whitelist/LLM

    def _build_whitelist(
        self, project_id: UUID, requirement_ids: list[UUID]
    ) -> tuple[WhitelistContext, list[dict[str, Any]]]:
        reqs = list(
            self.db.scalars(
                select(Requirement).where(
                    Requirement.project_id == project_id,
                    Requirement.id.in_(requirement_ids),
                )
            )
        )
        if len(reqs) != len(set(requirement_ids)):
            raise DraftValidationError("存在跨项目或不存在的 Requirement")

        matches = list(
            self.db.scalars(
                select(RequirementEvidenceMatch)
                .where(
                    RequirementEvidenceMatch.project_id == project_id,
                    RequirementEvidenceMatch.requirement_id.in_(requirement_ids),
                    RequirementEvidenceMatch.lifecycle_status == "active",
                )
                .options(
                    selectinload(RequirementEvidenceMatch.company_links).selectinload(
                        RequirementEvidenceMatchLink.document
                    ),
                    selectinload(RequirementEvidenceMatch.company_links).selectinload(
                        RequirementEvidenceMatchLink.chunk
                    ),
                    selectinload(RequirementEvidenceMatch.requirement).selectinload(
                        Requirement.evidence_links
                    ),
                )
            )
        )
        # Load tender evidence links
        tender_links = list(
            self.db.scalars(
                select(EvidenceLink)
                .join(Requirement, Requirement.id == EvidenceLink.requirement_id)
                .where(
                    Requirement.project_id == project_id,
                    EvidenceLink.requirement_id.in_(requirement_ids),
                )
                .options(
                    selectinload(EvidenceLink.document),
                    selectinload(EvidenceLink.chunk),
                )
            )
        )
        tender_by_req: dict[UUID, list[EvidenceLink]] = {}
        for el in tender_links:
            tender_by_req.setdefault(el.requirement_id, []).append(el)

        citations: dict[UUID, CitationMeta] = {}
        quotes: dict[str, CitationMeta] = {}
        source_rows: list[dict[str, Any]] = []
        req_status: dict[UUID, EvidenceMatchStatus] = {}
        req_match_id: dict[UUID, UUID] = {}
        gap_ids: set[UUID] = set()
        risk_ids: set[UUID] = set()
        scope_ids: set[UUID] = set()
        excluded_ids: set[UUID] = set()
        match_ids: set[UUID] = set()

        for match in matches:
            if match.review_status != MatchReviewStatus.confirmed:
                excluded_ids.add(match.requirement_id)
                continue
            match_ids.add(match.id)
            req_status[match.requirement_id] = match.status
            req_match_id[match.requirement_id] = match.id
            if match.status == EvidenceMatchStatus.insufficient_evidence:
                gap_ids.add(match.requirement_id)
            elif match.status == EvidenceMatchStatus.conflicting_evidence:
                risk_ids.add(match.requirement_id)
            elif match.status == EvidenceMatchStatus.not_applicable:
                scope_ids.add(match.requirement_id)

            for el in tender_by_req.get(match.requirement_id, []):
                quote = normalize_text(
                    (el.chunk.content if el.chunk else None) or el.notes or ""
                )
                qid = f"q_tender_{el.id.hex[:12]}"
                loc = _location_from_chunk(
                    document_file_name=el.document.file_name if el.document else None,
                    document_type=(
                        el.document.document_type.value if el.document else None
                    ),
                    chunk=el.chunk,
                    document_id=el.document_id,
                    chunk_id=el.chunk_id,
                    project_id=project_id,
                )
                meta = CitationMeta(
                    citation_id=el.id,
                    requirement_id=match.requirement_id,
                    match_id=match.id,
                    match_status=match.status,
                    source_role=ProposalDraftSourceRole.tender_requirement.value,
                    quote=quote,
                    quote_id=qid,
                    location=loc,
                )
                citations[el.id] = meta
                if quote:
                    quotes[qid] = meta
                source_rows.append(
                    {
                        "requirement_id": match.requirement_id,
                        "match_id": match.id,
                        "evidence_link_id": el.id,
                        "source_role": ProposalDraftSourceRole.tender_requirement.value,
                        "source_quote": quote or None,
                        "location_json": loc,
                    }
                )

            for clink in match.company_links:
                role = clink.role or "company_support"
                if role not in {
                    "company_support",
                    "company_conflict",
                    "company_scope_exclusion",
                    "requirement_scope",
                }:
                    role = "company_support"
                # Map requirement_scope to company_scope_exclusion for draft source enum
                source_role = (
                    ProposalDraftSourceRole.company_scope_exclusion.value
                    if role in {"requirement_scope", "company_scope_exclusion"}
                    else (
                        ProposalDraftSourceRole.company_conflict.value
                        if role == "company_conflict"
                        else ProposalDraftSourceRole.company_support.value
                    )
                )
                quote = normalize_text(clink.quote or "")
                qid = f"q_company_{clink.id.hex[:12]}"
                loc = _location_from_chunk(
                    document_file_name=clink.document.file_name if clink.document else None,
                    document_type=(
                        clink.document.document_type.value if clink.document else None
                    ),
                    chunk=clink.chunk,
                    document_id=clink.document_id,
                    chunk_id=clink.chunk_id,
                    project_id=project_id,
                )
                meta = CitationMeta(
                    citation_id=clink.id,
                    requirement_id=match.requirement_id,
                    match_id=match.id,
                    match_status=match.status,
                    source_role=source_role,
                    quote=quote,
                    quote_id=qid,
                    location=loc,
                )
                citations[clink.id] = meta
                if quote:
                    quotes[qid] = meta
                source_rows.append(
                    {
                        "requirement_id": match.requirement_id,
                        "match_id": match.id,
                        "evidence_link_id": clink.id,
                        "source_role": source_role,
                        "source_quote": quote or None,
                        "location_json": loc,
                    }
                )

        for rid in requirement_ids:
            if rid not in req_status and rid not in excluded_ids:
                excluded_ids.add(rid)

        whitelist = WhitelistContext(
            project_id=project_id,
            requirement_ids=set(requirement_ids),
            match_ids=match_ids,
            requirement_match_status=req_status,
            requirement_match_id=req_match_id,
            citation_ids=set(citations.keys()),
            citations=citations,
            quote_ids=set(quotes.keys()),
            quotes=quotes,
            gap_requirement_ids=gap_ids,
            risk_requirement_ids=risk_ids,
            scope_requirement_ids=scope_ids,
            excluded_requirement_ids=excluded_ids,
        )
        return whitelist, source_rows

    def _call_llm(
        self, run: ProposalDraftGenerationRun, whitelist: WhitelistContext
    ) -> dict[str, Any]:
        facts = []
        for rid, status_ in whitelist.requirement_match_status.items():
            cites = [
                {
                    "citation_id": str(c.citation_id),
                    "quote_id": c.quote_id,
                    "source_role": c.source_role,
                    "quote": c.quote[:500],
                    "location": {
                        k: c.location.get(k)
                        for k in (
                            "document_file_name",
                            "page_start",
                            "section",
                            "clause_id",
                        )
                    },
                }
                for c in whitelist.citations.values()
                if c.requirement_id == rid
            ]
            facts.append(
                {
                    "requirement_id": str(rid),
                    "match_id": str(whitelist.requirement_match_id[rid]),
                    "match_status": status_.value,
                    "expected_block_kind": {
                        EvidenceMatchStatus.supported: "supported_response",
                        EvidenceMatchStatus.partially_supported: "partial_response",
                        EvidenceMatchStatus.insufficient_evidence: "material_gap",
                        EvidenceMatchStatus.conflicting_evidence: "risk_item",
                        EvidenceMatchStatus.not_applicable: "scope_item",
                    }[status_],
                    "citations": cites,
                }
            )
        excluded = [
            {"requirement_id": str(rid), "reason": "not_confirmed_or_missing"}
            for rid in whitelist.excluded_requirement_ids
        ]
        user_payload = {
            "title": run.title,
            "mode": run.mode.value,
            "facts": facts,
            "excluded_requirements": excluded,
            "instructions": (
                "仅使用 facts 中的 citation_id 与 quote_id；"
                "不得填写或覆盖 location；不得编造 quote。"
            ),
        }
        result = self.llm.chat(
            [
                {"role": "system", "content": _SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": json.dumps(user_payload, ensure_ascii=False),
                },
            ],
            max_tokens=max(getattr(self.llm, "max_tokens", 1024) or 1024, 4096),
            temperature=0.1,
            request_id=str(run.id),
        )
        return _parse_llm_json(result.content)

    # ----------------------------------------------------- manual / review

    def create_manual_revision(
        self,
        project_id: UUID,
        draft_id: UUID,
        payload: ProposalDraftManualRevisionRequest,
        *,
        idempotency_key: str | None = None,
    ) -> ProposalDraftDetail:
        draft = self._lock_draft(project_id, draft_id)
        if draft.status == ProposalDraftStatus.reviewed:
            raise HTTPException(
                status_code=409,
                detail="已审核版本只读，请先 reopen 再编辑",
            )
        if draft.current_version_id is None:
            raise HTTPException(status_code=409, detail="草稿尚无当前版本")

        current = self.db.get(ProposalDraftVersion, draft.current_version_id)
        if current is None or current.project_id != project_id:
            raise HTTPException(status_code=409, detail="当前版本缺失")

        # Rebuild whitelist from current source snapshot requirement/match ids
        req_ids = sorted(
            {
                UUID(str(rid))
                for section in (payload.content_json.get("sections") or [])
                for block in (section.get("blocks") or [])
                for rid in (block.get("requirement_ids") or [])
            }
        )
        if not req_ids:
            # Fall back to snapshot sources
            sources = list(
                self.db.scalars(
                    select(ProposalDraftSource).where(
                        ProposalDraftSource.draft_version_id == current.id
                    )
                )
            )
            req_ids = sorted({s.requirement_id for s in sources if s.requirement_id})

        whitelist, source_rows = self._build_whitelist(project_id, req_ids)
        try:
            content = validate_structured_content(
                payload.content_json,
                whitelist,
                allow_manual_unreferenced=True,
            )
        except DraftValidationError as exc:
            raise HTTPException(status_code=422, detail=exc.message) from exc

        markdown = render_markdown(content)
        next_number = current.version_number + 1
        # Clear previous current flag
        current.is_current = False
        version = ProposalDraftVersion(
            project_id=project_id,
            draft_id=draft.id,
            parent_version_id=current.id,
            version_number=next_number,
            version_kind=ProposalDraftVersionKind.manual_revision,
            generation_run_id=None,
            source_snapshot_hash=current.source_snapshot_hash,
            content_json=content,
            content_markdown=markdown,
            created_by=payload.created_by,
            supersedes_version_id=current.id,
            is_current=True,
        )
        self.db.add(version)
        self.db.flush()
        for row in source_rows:
            self.db.add(
                ProposalDraftSource(
                    project_id=project_id,
                    draft_version_id=version.id,
                    requirement_id=row.get("requirement_id"),
                    match_id=row.get("match_id"),
                    evidence_link_id=row.get("evidence_link_id"),
                    source_role=ProposalDraftSourceRole(row["source_role"]),
                    source_quote=row.get("source_quote"),
                    location_json=row.get("location_json"),
                )
            )
        draft.current_version_id = version.id
        if draft.status == ProposalDraftStatus.reopened:
            draft.status = ProposalDraftStatus.draft_pending_review
        draft.updated_at = datetime.now(UTC)
        self.db.commit()
        return self.get_draft(project_id, draft_id)

    def mark_reviewed(
        self,
        project_id: UUID,
        draft_id: UUID,
        payload: ProposalDraftReviewRequest,
        *,
        idempotency_key: str | None = None,
    ) -> ProposalDraftDetail:
        draft = self._lock_draft(project_id, draft_id)
        if idempotency_key:
            existing = self.db.execute(
                select(ProposalDraftReview).where(
                    ProposalDraftReview.project_id == project_id,
                    ProposalDraftReview.draft_id == draft_id,
                    ProposalDraftReview.idempotency_key == idempotency_key,
                )
            ).scalar_one_or_none()
            if existing is not None:
                stored = existing.payload_hash
                expected = _payload_hash(
                    {
                        "action": payload.action.value,
                        "comment": payload.comment,
                        "actor_label": payload.actor_label,
                        "review_lock_version": payload.review_lock_version,
                    }
                )
                if stored and stored != expected:
                    raise HTTPException(
                        status_code=409,
                        detail="Idempotency-Key 已用于不同请求载荷",
                    )
                return self.get_draft(project_id, draft_id)

        if draft.review_lock_version != payload.review_lock_version:
            raise HTTPException(status_code=409, detail="草稿并发冲突，请刷新后重试")
        if draft.status == ProposalDraftStatus.reviewed:
            raise HTTPException(status_code=409, detail="草稿已审核")
        if draft.current_version_id is None:
            raise HTTPException(status_code=409, detail="无当前版本可审核")

        version = self.db.get(ProposalDraftVersion, draft.current_version_id)
        if version is None:
            raise HTTPException(status_code=409, detail="当前版本缺失")
        if content_has_unevidenced_manual(version.content_json or {}):
            raise HTTPException(
                status_code=422,
                detail="含「人工新增，尚未提供证据」的内容不可标记已审核",
            )

        review = ProposalDraftReview(
            project_id=project_id,
            draft_id=draft.id,
            draft_version_id=version.id,
            action=ProposalDraftReviewAction.mark_reviewed,
            comment=payload.comment,
            actor_label=payload.actor_label,
            actor_authn=ActorAuthn.unverified_local_operator,
            idempotency_key=idempotency_key,
            payload_hash=_payload_hash(
                {
                    "action": payload.action.value,
                    "comment": payload.comment,
                    "actor_label": payload.actor_label,
                    "review_lock_version": payload.review_lock_version,
                }
            ),
        )
        self.db.add(review)
        draft.status = ProposalDraftStatus.reviewed
        draft.review_lock_version += 1
        self.db.commit()
        return self.get_draft(project_id, draft_id)

    def reopen(
        self,
        project_id: UUID,
        draft_id: UUID,
        payload: ProposalDraftReopenRequest,
        *,
        idempotency_key: str | None = None,
    ) -> ProposalDraftDetail:
        draft = self._lock_draft(project_id, draft_id)
        if idempotency_key:
            existing = self.db.execute(
                select(ProposalDraftReview).where(
                    ProposalDraftReview.project_id == project_id,
                    ProposalDraftReview.draft_id == draft_id,
                    ProposalDraftReview.idempotency_key == idempotency_key,
                )
            ).scalar_one_or_none()
            if existing is not None:
                return self.get_draft(project_id, draft_id)

        if draft.review_lock_version != payload.review_lock_version:
            raise HTTPException(status_code=409, detail="草稿并发冲突，请刷新后重试")
        if draft.status != ProposalDraftStatus.reviewed:
            raise HTTPException(status_code=409, detail="仅已审核草稿可 reopen")
        if draft.current_version_id is None:
            raise HTTPException(status_code=409, detail="无当前版本")

        review = ProposalDraftReview(
            project_id=project_id,
            draft_id=draft.id,
            draft_version_id=draft.current_version_id,
            action=ProposalDraftReviewAction.reopen,
            comment=payload.comment,
            actor_label=payload.actor_label,
            actor_authn=ActorAuthn.unverified_local_operator,
            idempotency_key=idempotency_key,
            payload_hash=_payload_hash(
                {
                    "action": "reopen",
                    "comment": payload.comment,
                    "actor_label": payload.actor_label,
                    "review_lock_version": payload.review_lock_version,
                }
            ),
        )
        self.db.add(review)
        draft.status = ProposalDraftStatus.reopened
        draft.review_lock_version += 1
        self.db.commit()
        return self.get_draft(project_id, draft_id)

    # ----------------------------------------------------------------- export

    def export(
        self, project_id: UUID, draft_id: UUID, *, fmt: str
    ) -> tuple[bytes, str, str]:
        """Return (body, media_type, filename)."""
        draft = self.get_draft(project_id, draft_id)
        if draft.status != ProposalDraftStatus.reviewed:
            raise HTTPException(status_code=409, detail="仅已审核当前版本可导出")
        if draft.current_version is None:
            raise HTTPException(status_code=409, detail="无当前版本")
        if draft.has_unevidenced_manual_content:
            raise HTTPException(
                status_code=422,
                detail="含尚未提供证据的人工内容不可导出",
            )
        content = draft.current_version.content_json
        markdown = draft.current_version.content_markdown or render_markdown(content)
        safe_title = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", draft.title)[:80] or "draft"
        if fmt == "markdown":
            return (
                markdown.encode("utf-8"),
                "text/markdown; charset=utf-8",
                f"{safe_title}.md",
            )
        if fmt == "docx":
            body = self._render_docx(draft.title, content, markdown)
            return (
                body,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{safe_title}.docx",
            )
        raise HTTPException(status_code=422, detail="format 仅支持 markdown|docx")

    def _render_docx(
        self, title: str, content: dict[str, Any], markdown: str
    ) -> bytes:
        try:
            from docx import Document as DocxDocument
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise HTTPException(
                status_code=501,
                detail="python-docx 未安装，无法导出 DOCX",
            ) from exc

        doc = DocxDocument()
        cover = doc.add_paragraph(DISCLAIMER)
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(normalize_text(title, max_len=512), level=0)
        pending = doc.add_paragraph("待人工复核，未提交")
        pending.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for section in content.get("sections") or []:
            doc.add_heading(normalize_text(section.get("title"), max_len=256), level=1)
            for block in section.get("blocks") or []:
                doc.add_heading(f"[{block.get('block_kind')}]", level=2)
                doc.add_paragraph(normalize_text(block.get("content")))
                cites = block.get("citation_ids") or []
                if cites:
                    doc.add_paragraph(f"引用: {', '.join(str(c) for c in cites)}")
                for loc in block.get("locations") or []:
                    parts = [
                        str(loc[k])
                        for k in (
                            "document_file_name",
                            "page_start",
                            "section",
                            "clause_id",
                        )
                        if loc.get(k) is not None
                    ]
                    if parts:
                        doc.add_paragraph("定位: " + " / ".join(parts))

        matrix = content.get("compliance_matrix") or []
        if matrix:
            doc.add_heading("合规准备矩阵", level=1)
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Requirement"
            hdr[1].text = "Disposition"
            hdr[2].text = "Citations"
            for row in matrix:
                cells = table.add_row().cells
                cells[0].text = str(row.get("requirement_id"))
                cells[1].text = str(row.get("disposition"))
                cells[2].text = ", ".join(row.get("citation_ids") or [])

        warnings = content.get("warnings") or []
        if warnings:
            doc.add_heading("风险与待核验", level=1)
            for w in warnings:
                doc.add_paragraph(
                    f"[{w.get('warning_type')}] {w.get('requirement_id')}: "
                    f"{normalize_text(w.get('content'))}"
                )

        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = DISCLAIMER

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # --------------------------------------------------------------- helpers

    def _to_summary(self, draft: ProposalDraft) -> ProposalDraftSummary:
        version = None
        version_number = None
        has_unevidenced = False
        counts = {
            "eligible_requirement_count": 0,
            "material_gap_count": 0,
            "risk_count": 0,
            "scope_count": 0,
        }
        if draft.current_version_id:
            version = self.db.get(ProposalDraftVersion, draft.current_version_id)
            if version:
                version_number = version.version_number
                has_unevidenced = content_has_unevidenced_manual(
                    version.content_json or {}
                )
                counts = self._count_from_content(version.content_json or {})

        last_reviewed = None
        if draft.status == ProposalDraftStatus.reviewed:
            review = self.db.execute(
                select(ProposalDraftReview)
                .where(
                    ProposalDraftReview.draft_id == draft.id,
                    ProposalDraftReview.action
                    == ProposalDraftReviewAction.mark_reviewed,
                )
                .order_by(ProposalDraftReview.created_at.desc())
                .limit(1)
            ).scalar_one_or_none()
            if review:
                last_reviewed = review.created_at

        export_allowed = (
            draft.status == ProposalDraftStatus.reviewed and not has_unevidenced
        )
        return ProposalDraftSummary(
            id=draft.id,
            project_id=draft.project_id,
            title=draft.title,
            status=draft.status,
            current_version_id=draft.current_version_id,
            current_version_number=version_number,
            created_by=draft.created_by,
            review_lock_version=draft.review_lock_version,
            created_at=draft.created_at,
            updated_at=draft.updated_at,
            last_reviewed_at=last_reviewed,
            has_unevidenced_manual_content=has_unevidenced,
            export_allowed=export_allowed,
            eligible_requirement_count=counts["eligible_requirement_count"],
            material_gap_count=counts["material_gap_count"],
            risk_count=counts["risk_count"],
            scope_count=counts["scope_count"],
        )

    def _count_from_content(self, content: dict[str, Any]) -> dict[str, int]:
        matrix = content.get("compliance_matrix") or []
        eligible = 0
        gaps = 0
        risks = 0
        scopes = 0
        for row in matrix:
            disp = row.get("disposition")
            if disp in {"responded", "partially_responded"}:
                eligible += 1
            elif disp == "material_gap":
                gaps += 1
            elif disp == "risk_review":
                risks += 1
            elif disp == "scope_review":
                scopes += 1
        # Also count blocks if matrix empty
        if not matrix:
            for section in content.get("sections") or []:
                for block in section.get("blocks") or []:
                    kind = block.get("block_kind")
                    if kind in {"supported_response", "partial_response"}:
                        eligible += 1
                    elif kind == "material_gap":
                        gaps += 1
                    elif kind == "risk_item":
                        risks += 1
                    elif kind == "scope_item":
                        scopes += 1
        return {
            "eligible_requirement_count": eligible,
            "material_gap_count": gaps,
            "risk_count": risks,
            "scope_count": scopes,
        }

    def _version_summary(self, version: ProposalDraftVersion) -> ProposalDraftVersionSummary:
        return ProposalDraftVersionSummary(
            id=version.id,
            project_id=version.project_id,
            draft_id=version.draft_id,
            parent_version_id=version.parent_version_id,
            version_number=version.version_number,
            version_kind=version.version_kind,
            generation_run_id=version.generation_run_id,
            source_snapshot_hash=version.source_snapshot_hash,
            created_by=version.created_by,
            supersedes_version_id=version.supersedes_version_id,
            is_current=version.is_current,
            created_at=version.created_at,
            has_unevidenced_manual_content=content_has_unevidenced_manual(
                version.content_json or {}
            ),
        )

    def _version_detail(
        self, project_id: UUID, version_id: UUID
    ) -> ProposalDraftVersionDetail:
        version = self.db.execute(
            select(ProposalDraftVersion)
            .where(
                ProposalDraftVersion.id == version_id,
                ProposalDraftVersion.project_id == project_id,
            )
            .options(selectinload(ProposalDraftVersion.sources))
        ).scalar_one_or_none()
        if version is None:
            raise HTTPException(status_code=404, detail="版本不存在")
        return self._version_detail_from_model(version)

    def _version_detail_from_model(
        self, version: ProposalDraftVersion
    ) -> ProposalDraftVersionDetail:
        summary = self._version_summary(version)
        sources = [
            ProposalDraftSourceRead.model_validate(s) for s in (version.sources or [])
        ]
        return ProposalDraftVersionDetail(
            **summary.model_dump(),
            content_json=version.content_json,
            content_markdown=version.content_markdown,
            sources=sources,
        )
